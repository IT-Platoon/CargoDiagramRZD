import base64
import tempfile

from docx import Document


async def create_document(
    result: dict,
    docxfile: tempfile.NamedTemporaryFile,
) -> None:
    doc = Document()
    doc.add_heading('Расчетно-пояснительная записка')
    paragraph = doc.add_paragraph(f"Продольное смещение грузов в вагоне: l(c) = {result['longitudinal_displacement_in_car']} мм. ")
    paragraph.add_run(result['longitudinal_displacement_flag']).bold = True
    doc.add_paragraph(f"Продольное смещение грузов с вагоном: l(c) = {result['longitudinal_displacement_with_car']} мм")
    doc.add_paragraph(f"Поперечное смещение в вагоне: b = {result['lateral_displacement_in_car']} мм")
    doc.add_paragraph(f"Поперечное смещение с вагоном: b = {result['lateral_displacement_with_car']} мм")
    doc.add_paragraph(f"Высота центра тяжести в вагоне: H(цт О) = {result['height_center_gravity_in_car']} мм")
    doc.add_paragraph(f"Общая высота ЦТ: H(цт) = {result['general_height_center_gravity']} мм")
    doc.add_paragraph(f"Расчет наветренной поверхности: S(bok) = {result['windward_surface']} м2")
    doc.add_paragraph(f"Удельная продольная инерционная сила на одну тонну веса груза: а(пр) = {result['specific_length_inertial_force_per_ton_cargo_weight']} тс/с")
    for index, cargo in enumerate(result['cargo'], start=1):
        doc.add_heading(f'Груз {index}', 3)
        doc.add_paragraph(f"Длина {cargo['length']}", style='List Bullet 2')
        doc.add_paragraph(f"Ширина {cargo['width']}", style='List Bullet 2')
        doc.add_paragraph(f"Высота {cargo['height']}", style='List Bullet 2')
        doc.add_paragraph(f"Масса {cargo['weight']}", style='List Bullet 2')
        doc.add_paragraph(f"Продольная инерционная сила: F(пр) = {cargo['longitudinal_inertial_force']} тс", style='List Bullet 2')
        doc.add_paragraph(f"Ветровая нагрузка: W(е) = {cargo['wind_load']} тс", style='List Bullet 2')
        doc.add_paragraph(f"Сила трения в продольном направлении: F(пр тр) = {cargo['friction_force_longitudinal_direction']} тс", style='List Bullet 2')
    doc.save(docxfile.name)


async def get_report(result: dict) -> str:
    with tempfile.NamedTemporaryFile(mode="rb+", suffix=".docx") as docxfile:
        await create_document(result, docxfile)
        docxfile.seek(0)
        report = base64.b64encode(docxfile.read()).decode()
    return report
